"""
Document Generator module for AutoDoc Agent.
Converts a content dictionary into a formatted .docx Word document using python-docx.
"""

import os
import re
import logging
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import config
from models import SectionContent

logger = logging.getLogger(__name__)


# ── Styling Constants ───────────────────────────────────────────────────────

FONT_NAME = "Calibri"
TITLE_SIZE = Pt(24)
HEADING_SIZE = Pt(16)
BODY_SIZE = Pt(11)
SMALL_SIZE = Pt(9)

COLOR_PRIMARY = RGBColor(0x1A, 0x1A, 0x2E)     # Dark navy
COLOR_HEADING = RGBColor(0x2D, 0x3A, 0x8C)      # Professional blue
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)          # Dark gray
COLOR_LIGHT = RGBColor(0x88, 0x88, 0x88)         # Light gray for metadata
COLOR_TABLE_HEADER_BG = RGBColor(0x2D, 0x3A, 0x8C)  # Blue header
COLOR_TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)  # White text on header
COLOR_TABLE_ALT_ROW = RGBColor(0xF2, 0xF4, 0xF8)    # Light blue-gray alternate rows


# ── Helper Functions ────────────────────────────────────────────────────────

def _set_run_font(run, size=BODY_SIZE, color=COLOR_BODY, bold=False):
    """Apply consistent font styling to a text run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold


def _add_styled_paragraph(doc, text, size=BODY_SIZE, color=COLOR_BODY,
                           bold=False, alignment=None, space_after=Pt(6)):
    """Add a paragraph with consistent styling."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, size=size, color=color, bold=bold)
    if alignment:
        para.alignment = alignment
    para.paragraph_format.space_after = space_after
    return para


def _set_cell_shading(cell, color: RGBColor):
    """Set the background shading of a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _sanitize_filename(text: str) -> str:
    """Convert text to a safe filename slug."""
    slug = text.lower().strip()
    slug = re.sub(r"[^a-z0-9]+", "_", slug)
    slug = slug.strip("_")
    return slug[:50]  # cap length


# ── Table Builder ───────────────────────────────────────────────────────────

def _add_table(doc, table_data: list[dict]):
    """
    Add a professionally styled table to the document.
    
    Args:
        doc: The Document object.
        table_data: List of dictionaries, each dict is one row.
    """
    if not table_data:
        return

    headers = list(table_data[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # ── Header row ──────────────────────────────────────────────────────
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(str(header))
        _set_run_font(run, size=Pt(10), color=COLOR_TABLE_HEADER_FG, bold=True)
        _set_cell_shading(cell, COLOR_TABLE_HEADER_BG)

    # ── Data rows ───────────────────────────────────────────────────────
    for row_idx, row_data in enumerate(table_data):
        row = table.add_row()
        for col_idx, header in enumerate(headers):
            cell = row.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            value = row_data.get(header, "")
            run = para.add_run(str(value))
            _set_run_font(run, size=Pt(10), color=COLOR_BODY)

            # Alternate row shading
            if row_idx % 2 == 0:
                _set_cell_shading(cell, COLOR_TABLE_ALT_ROW)

    # Add spacing after table
    doc.add_paragraph()


# ── Public API ──────────────────────────────────────────────────────────────

def generate_document(
    document_type: str,
    sections: dict[str, SectionContent],
) -> str:
    """
    Generate a formatted .docx Word document from section content.
    
    Args:
        document_type: The type of document (used as title).
        sections: Dictionary mapping section titles to their content.
        
    Returns:
        The file path of the generated .docx file.
    """
    # Ensure output directory exists
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)

    doc = Document()

    # ── Document Margins ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Title ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(document_type)
    _set_run_font(title_run, size=TITLE_SIZE, color=COLOR_PRIMARY, bold=True)
    title_para.paragraph_format.space_after = Pt(4)

    # ── Horizontal rule (subtle line under title) ───────────────────────
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run("─" * 60)
    _set_run_font(line_run, size=SMALL_SIZE, color=COLOR_LIGHT)
    line_para.paragraph_format.space_after = Pt(4)

    # ── Metadata ────────────────────────────────────────────────────────
    now = datetime.now()
    meta_text = (
        f"Generated on {now.strftime('%B %d, %Y at %I:%M %p')}  •  "
        f"AutoDoc Agent"
    )
    _add_styled_paragraph(
        doc, meta_text,
        size=SMALL_SIZE, color=COLOR_LIGHT,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(20),
    )

    # ── Sections ────────────────────────────────────────────────────────
    for section_title, section_content in sections.items():
        # Section heading
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(section_content.title)
        _set_run_font(heading_run, size=HEADING_SIZE, color=COLOR_HEADING, bold=True)
        heading_para.paragraph_format.space_before = Pt(16)
        heading_para.paragraph_format.space_after = Pt(8)

        # Section body text
        if section_content.content:
            # Split content into paragraphs (respect double newlines)
            paragraphs = section_content.content.split("\n\n")
            for para_text in paragraphs:
                clean_text = para_text.strip()
                if clean_text:
                    _add_styled_paragraph(doc, clean_text, space_after=Pt(8))

        # Section table (if any)
        if section_content.table_data:
            _add_table(doc, section_content.table_data)

    # ── Footer note ─────────────────────────────────────────────────────
    doc.add_paragraph()  # spacing
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "This document was autonomously generated by AutoDoc Agent. "
        "Please review and verify all content before distribution."
    )
    _set_run_font(footer_run, size=SMALL_SIZE, color=COLOR_LIGHT)

    # ── Save ────────────────────────────────────────────────────────────
    slug = _sanitize_filename(document_type)
    date_str = now.strftime("%Y-%m-%d")
    filename = f"{slug}_{date_str}.docx"
    filepath = os.path.join(config.OUTPUT_DIR, filename)

    doc.save(filepath)
    logger.info(f"Document saved to: {filepath}")

    return filepath
